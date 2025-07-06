from django.shortcuts import render,redirect,get_object_or_404
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login,authenticate
from django.contrib.auth.decorators import login_required
from .forms import SignupForm,RoomForm
from .models import Room,RoomContent,Document
from django.contrib import messages
from django.http import JsonResponse
from django.contrib.auth.models import User
from django.contrib.auth.backends import ModelBackend
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from django.utils import timezone
from .models import UserStatus
from django.contrib.auth import logout
from django.contrib.sessions.models import Session
from django.contrib.auth.forms import AuthenticationForm
from django.views.decorators.csrf import csrf_exempt
from django.utils.timezone import now
from django.core.cache import cache
from django.http import HttpResponse
from django.core.mail import send_mail
from django.conf import settings
import random
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from io import BytesIO
from django.http import HttpResponse
from docx import Document
from django.views.decorators.csrf import csrf_exempt
from django.core.mail import send_mail
from django.http import JsonResponse
import json
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Room, RoomContent
from django.contrib.auth.models import User
from django.core.mail import send_mail
from django.conf import settings
from django.contrib.auth import login
from django.urls import reverse

@login_required
def document_view(request, room_id):
    room = get_object_or_404(Room, id=room_id)
    content, created = RoomContent.objects.get_or_create(room=room)
    return render(request, 'document.html', {
        'room_id': room.id,
        'room_name': room.name,
        'content': content.content,
        'user': request.user,
    })

@login_required
def invite_user(request, room_id):
    if request.method == 'POST':
        email = request.POST['email']
        room = get_object_or_404(Room, id=room_id)
        invite_link = request.build_absolute_uri(reverse('join_document', args=[room.id]))
        
        send_mail(
            subject='Collaborative Document Invite',
            message=f"You've been invited to join the room '{room.name}'. Click here: {invite_link}",
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=[email],
        )
        return redirect('document', room_id=room.id)

    return redirect('document', room_id=room_id)

def join_document(request, room_id):
    if not request.user.is_authenticated:
        request.session['pending_room'] = room_id
        return redirect('login')

    return redirect('document', room_id=room_id)
@csrf_exempt
def share_document(request, room_id):
    if request.method == "POST":
        data = json.loads(request.body)
        emails = data.get("emails", [])
        message = data.get("message", "")
        room_name = data.get("room_name", "")
        room_url = data.get("room_url", "")

        subject = f"Invitation to collaborate on: {room_name}"
        full_message = f"{message}\n\nClick to join: {room_url}"

        try:
            send_mail(subject, full_message, 'your_email@example.com', emails)
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})

    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})


def export_docx(request, room_id):
    from .models import RoomContent
    content = RoomContent.objects.get(room_id=room_id).content

    doc = Document()
    doc.add_heading("Exported Document", 0)
    for line in content.splitlines():
        doc.add_paragraph(line)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=document.docx'
    doc.save(response)
    return response

def export_pdf(request, room_id):
    # Load content from DB
    from .models import RoomContent
    content = RoomContent.objects.get(room_id=room_id).content

    # Create PDF
    buffer = BytesIO()
    p = canvas.Canvas(buffer)
    p.drawString(100, 800, "Exported Document")
    text_object = p.beginText(100, 780)
    for line in content.splitlines():
        text_object.textLine(line)
    p.drawText(text_object)
    p.showPage()
    p.save()

    buffer.seek(0)
    return HttpResponse(buffer, content_type='application/pdf')
@csrf_exempt
def update_editing_status(request, room_id):
    if request.method == "POST" and request.user.is_authenticated:
        cache_key = f"editing:{room_id}:{request.user.username}"
        cache.set(cache_key, True, timeout=10)  # user is "editing" for 10 seconds
        return JsonResponse({'status': 'ok'})
    return JsonResponse({'status': 'error'}, status=400)

# Get users who are currently editing
def get_editing_users(request, room_id):
    room = Room.objects.get(id=room_id)
    users = room.participants.all().values_list('username', flat=True)
    editing_users = []

    for username in users:
        cache_key = f"editing:{room_id}:{username}"
        if cache.get(cache_key):
            editing_users.append(username)

    return JsonResponse({'editing': editing_users})

def home(request):
    return render(request,"home.html")
def register(request):
    if request.method == 'POST':
        username = request.POST['username']
        email = request.POST['email']
        password1 = request.POST['password1']
        password2 = request.POST['password2']

        if password1 == password2:
            if User.objects.filter(username=username).exists():
                messages.error(request, 'Username already exists.')
            elif User.objects.filter(email=email).exists():
                messages.error(request, 'Email already registered.')
            else:
                user = User.objects.create_user(username=username, email=email, password=password1)
                user.save()

                # Authenticate the user
                user = authenticate(request, username=username, password=password1)
                if user is not None:
                    login(request, user)

                return redirect('dashboard')  # or wherever you want
        else:
            messages.error(request, 'Passwords do not match.')
    return render(request, 'register.html')
@login_required
def dashboard(request):
    rooms = Room.objects.all()  # Get all rooms
    return render(request, "dashboard.html", {"rooms": rooms})
@login_required
def create_room(request):
    if request.method == "POST":
        room_name = request.POST.get("room_name")
        room_password = request.POST.get("room_password")

        if Room.objects.filter(name=room_name).exists():
            return render(request, "create_room.html", {"error": "Room name already exists!"})

        Room.objects.create(name=room_name, password=room_password, admin=request.user)
        return redirect("dashboard")

    return render(request, "create_room.html")
@login_required
def join_room(request, room_id):
    room = Room.objects.get(id=room_id)

    if request.method == "POST":
        entered_password = request.POST.get("password")
        if entered_password == room.password:
            room.participants.add(request.user)

            # ✅ Create or update UserStatus with valid room
            UserStatus.objects.update_or_create(
                user=request.user,
                room=room,
                defaults={'is_logged_in': True}
            )

            return redirect("room", room_id=room.id)
        else:
            messages.error(request, "Incorrect password. Try again.")
    
    return render(request, "join_room.html", {"room": room})
@login_required
def room_view(request, room_id):
    room = get_object_or_404(Room, id=room_id)
    content, _ = RoomContent.objects.get_or_create(room=room)
    return render(request, "document.html", {
        "room": room,
        "room_id": room_id,  # ✅ Required for export links
        "content": content.content
    })
@csrf_exempt
@login_required
def save_document(request, room_id):
    if request.method == "POST":
        data = json.loads(request.body)
        content = data.get("content", "")
        room = get_object_or_404(Room, id=room_id)

        room_content, created = RoomContent.objects.get_or_create(room=room)
        room_content.content = content
        room_content.save()

        return JsonResponse({"status": "success"})
def get_document(request, room_id):
    room = get_object_or_404(Room, id=room_id)
    room_content, created = RoomContent.objects.get_or_create(room=room)
    return JsonResponse({'content': room_content.content})

def get_active_users(request, room_id):
    sessions = Session.objects.filter(expire_date__gte=now())
    active_user_ids = []

    for session in sessions:
        data = session.get_decoded()
        user_id = data.get('_auth_user_id')
        if user_id:
            active_user_ids.append(int(user_id))

    users_in_room = UserStatus.objects.filter(room_id=room_id, user_id__in=active_user_ids)
    # Only show users who are really logged in
    return JsonResponse({
        "users": [u.user.username for u in users_in_room]
    })
def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()

            # Check for active sessions of the same user
            sessions = Session.objects.filter(expire_date__gte=timezone.now())
            for session in sessions:
                data = session.get_decoded()
                if str(user.id) == str(data.get('_auth_user_id')):
                    return render(request, 'login.html', {
                        'form': form,
                        'message': "You're already logged in from another device."
                    })

            login(request, user)

            # ✅ DO NOT update UserStatus here — wait until the user joins a room
            return redirect('dashboard')
    else:
        form = AuthenticationForm()

    return render(request, 'login.html', {'form': form})

def custom_logout(request):
    try:
      if request.user.is_authenticated:
        UserStatus.objects.filter(user=request.user).delete() 
    except UserStatus.DoesNotExist:
            pass
    logout(request)
    return redirect('home')
def logout_room_users(request, room_id):
    if request.user.is_superuser:  # Only admin allowed
        count, _ = UserStatus.objects.filter(room_id=room_id).delete()
        return HttpResponse(f"{count} users logged out from room {room_id}")
    return HttpResponse("Unauthorized", status=401)
def forgot_password(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        try:
            user = User.objects.get(email=email)
            otp = str(random.randint(100000, 999999))
            request.session['reset_email'] = email
            request.session['otp'] = otp
            send_mail(
                'Password Reset OTP',
                f'Your OTP for password reset is: {otp}',
                settings.EMAIL_HOST_USER,
                [email],
                fail_silently=False,
            )
            return redirect('verify_otp')
        except User.DoesNotExist:
            messages.error(request, "No account with this email.")
            return redirect('forgot_password')
    return render(request, 'forgot_password.html')

def verify_otp(request):
    if request.method == 'POST':
        otp_entered = request.POST.get('otp')
        otp_session = request.session.get('otp')
        if otp_entered == otp_session:
            return redirect('reset_password')
        else:
            messages.error(request, "Invalid OTP")
            return redirect('verify_otp')
    return render(request, 'verify_otp.html')

def reset_password(request):
    if request.method == 'POST':
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')
        if password == confirm_password:
            email = request.session.get('reset_email')
            user = User.objects.get(email=email)
            user.set_password(password)
            user.save()
            messages.success(request, "Password reset successful. Please login.")
            return redirect('login')  # Update with your login URL name
        else:
            messages.error(request, "Passwords do not match.")
            return redirect('reset_password')
    return render(request, 'reset_password.html')